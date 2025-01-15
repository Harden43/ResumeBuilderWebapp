from flask import Flask, render_template, request, send_file, session
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.colors import HexColor

app = Flask(__name__)
app.secret_key = 'momislove'  # Replace with a secure random string

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        resume_data = {
            'name': request.form.get('name', ''),
            'address': request.form.get('address', ''),
            'email': request.form.get('email', ''),
            'phone': request.form.get('phone', ''),
            'linkedin': request.form.get('linkedin', ''),
            'summary': request.form.get('summary', ''),
            'experience': [],
            'education': [],
            'skills': request.form.get('skills', '').split(','),
            'certifications': request.form.get('certifications', '').split(','),
            'volunteer': request.form.get('volunteer', ''),
            'projects': request.form.get('projects', '')
        }
        
        # Process dynamic job experience fields
        job_titles = request.form.getlist('job_title')
        companies = request.form.getlist('company')
        dates = request.form.getlist('dates')
        responsibilities = request.form.getlist('responsibilities')
        
        for i in range(len(job_titles)):
            if job_titles[i] or companies[i] or dates[i] or responsibilities[i]:
                resume_data['experience'].append({
                    'title': job_titles[i],
                    'company': companies[i],
                    'dates': dates[i],
                    'responsibilities': responsibilities[i].split('|') if responsibilities[i] else []
                })
        
        # Process dynamic education fields
        degrees = request.form.getlist('degree')
        institutions = request.form.getlist('institution')
        graduation_dates = request.form.getlist('graduation_date')
        
        for i in range(len(degrees)):
            if degrees[i] or institutions[i] or graduation_dates[i]:
                resume_data['education'].append({
                    'degree': degrees[i],
                    'institution': institutions[i],
                    'graduation_date': graduation_dates[i]
                })
        
        # Store resume data in session
        session['resume_data'] = resume_data
        
        return render_template('resume.html', resume=resume_data)
    
    # Retrieve stored data from session if available
    resume_data = session.get('resume_data', {})
    return render_template('index.html', resume=resume_data)

@app.route('/download/<format>')
def download(format):
    resume_data = session.get('resume_data', {})

    if format == 'pdf':
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Modify existing styles
        styles['Heading1'].fontSize = 16
        styles['Heading1'].textColor = HexColor('#4A4A4A')
        styles['Heading2'].fontSize = 14
        styles['Heading2'].textColor = HexColor('#4A4A4A')
        styles['Normal'].fontSize = 10
        styles['Normal'].textColor = HexColor('#333333')
        
        story = []
        
        # Name
        story.append(Paragraph(resume_data['name'], styles['Heading1']))
        
        # Contact Info
        contact_info = f"{resume_data['address']} | {resume_data['email']} | {resume_data['phone']}"
        if resume_data['linkedin']:
            contact_info += f" | {resume_data['linkedin']}"
        story.append(Paragraph(contact_info, styles['Normal']))
        
        # Summary
        if resume_data['summary']:
            story.append(Paragraph('Summary', styles['Heading2']))
            story.append(Paragraph(resume_data['summary'], styles['Normal']))
        
        # Experience
        if resume_data['experience']:
            story.append(Paragraph('Experience', styles['Heading2']))
            for job in resume_data['experience']:
                story.append(Paragraph(f"<b>{job['title']}</b>", styles['Normal']))
                story.append(Paragraph(f"{job['company']} | {job['dates']}", styles['Normal']))
                for resp in job['responsibilities']:
                    story.append(Paragraph(f"• {resp}", styles['Normal']))
        
        # Education
        if resume_data['education']:
            story.append(Paragraph('Education', styles['Heading2']))
            for edu in resume_data['education']:
                story.append(Paragraph(f"<b>{edu['degree']}</b>", styles['Normal']))
                story.append(Paragraph(f"{edu['institution']} | {edu['graduation_date']}", styles['Normal']))
        
        # Skills
        if resume_data['skills']:
            story.append(Paragraph('Skills', styles['Heading2']))
            story.append(Paragraph(', '.join(resume_data['skills']), styles['Normal']))
        
        # Certifications
        if resume_data['certifications']:
            story.append(Paragraph('Certifications', styles['Heading2']))
            for cert in resume_data['certifications']:
                story.append(Paragraph(f"• {cert}", styles['Normal']))
        
        # Volunteer Experience
        if resume_data['volunteer']:
            story.append(Paragraph('Volunteer Experience', styles['Heading2']))
            story.append(Paragraph(resume_data['volunteer'], styles['Normal']))
        
        # Projects
        if resume_data['projects']:
            story.append(Paragraph('Projects', styles['Heading2']))
            story.append(Paragraph(resume_data['projects'], styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return send_file(buffer, 
                         as_attachment=True, 
                         download_name='resume.pdf', 
                         mimetype='application/pdf')
    
    elif format == 'docx':
        doc = Document()
        
        # Styles
        styles = doc.styles
        style_heading1 = styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
        style_heading1.font.size = Pt(16)
        style_heading1.font.color.rgb = RGBColor(74, 74, 74)
        
        style_heading2 = styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
        style_heading2.font.size = Pt(14)
        style_heading2.font.color.rgb = RGBColor(74, 74, 74)
        
        # Name
        doc.add_paragraph(resume_data['name'], style='Heading1')
        
        # Contact Info
        contact_info = f"{resume_data['address']} | {resume_data['email']} | {resume_data['phone']}"
        if resume_data['linkedin']:
            contact_info += f" | {resume_data['linkedin']}"
        doc.add_paragraph(contact_info)
        
        # Summary
        if resume_data['summary']:
            doc.add_paragraph('Summary', style='Heading2')
            doc.add_paragraph(resume_data['summary'])
        
        # Experience
        if resume_data['experience']:
            doc.add_paragraph('Experience', style='Heading2')
            for job in resume_data['experience']:
                p = doc.add_paragraph()
                p.add_run(job['title']).bold = True
                doc.add_paragraph(f"{job['company']} | {job['dates']}")
                for resp in job['responsibilities']:
                    doc.add_paragraph(resp, style='List Bullet')
        
        # Education
        if resume_data['education']:
            doc.add_paragraph('Education', style='Heading2')
            for edu in resume_data['education']:
                p = doc.add_paragraph()
                p.add_run(edu['degree']).bold = True
                doc.add_paragraph(f"{edu['institution']} | {edu['graduation_date']}")
        
        # Skills
        if resume_data['skills']:
            doc.add_paragraph('Skills', style='Heading2')
            doc.add_paragraph(', '.join(resume_data['skills']))
        
        # Certifications
        if resume_data['certifications']:
            doc.add_paragraph('Certifications', style='Heading2')
            for cert in resume_data['certifications']:
                doc.add_paragraph(cert, style='List Bullet')
        
        # Volunteer Experience
        if resume_data['volunteer']:
            doc.add_paragraph('Volunteer Experience', style='Heading2')
            doc.add_paragraph(resume_data['volunteer'])
        
        # Projects
        if resume_data['projects']:
            doc.add_paragraph('Projects', style='Heading2')
            doc.add_paragraph(resume_data['projects'])
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer,
                         as_attachment=True,
                         download_name="resume.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    else:
        return "Unsupported format", 400

if __name__ == '__main__':
    # app.run(debug=True)
    port = int(os.environ.get('PORT', 5000))  # Default to 5000 if PORT is not set
    app.run(host='0.0.0.0', port=port)
